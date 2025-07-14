from .exceptions.parser_error import ParserError
from parser_service.app.parsers.parser import Parser
import docx
import logging

logger = logging.getLogger(__name__)


class DocxParser(Parser):
    def extract_text(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text = [p.text for p in doc.paragraphs if p.text]

            # Extract tables
            try:
                for table in doc.tables:
                    for row_index, row in enumerate(table.rows):
                        try:
                            row_text = [cell.text.strip() for cell in row.cells if
                                cell.text.strip()]
                            if row_text:
                                text.append(" | ".join(row_text))
                        except Exception as e:
                            logger.warning(
                                f"Failed to extract row {row_index + 1} in "
                                f"a table from '{file_path}': {e}"
                            )
            except Exception as e:
                logger.warning(
                    f"Error extracting tables from '{file_path}': {e}"
                )

            full_text = "\n".join(text)

            if not full_text.strip():
                logger.error(
                    f"DOCX file '{file_path}' parsed but contained no extractable text.")
                raise ParserError("DOCX file contains no extractable text.")
            return full_text
        except FileNotFoundError as e:
            logger.error(f"DOCX file '{file_path}' not found.")
            raise ParserError("DOCX file not found.") from e
        except PermissionError as e:
            logger.error(f"No permission to read DOCX file: '{file_path}'.")
            raise  ParserError("No permission to read DOCX file.") from e
        except IsADirectoryError as e:
            logger.error(f"Expected a file but got a directory: '{file_path}'.")
            raise ParserError("Expected a file but got a directory.") from e
        except Exception as e:
            logger.critical(
                f"Unexpected error while parsing DOCX file '{file_path}': {e}",
                exc_info=True)
            raise ParserError(
                "Unknown error occurred during DOCX parsing.") from e
